#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate an independently constructed synthetic schedule DOCX."""

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_lines(cell, lines):
    cell.text = ""
    for index, line in enumerate(lines):
        if index > 0:
            cell.add_paragraph()
        paragraph = cell.paragraphs[index]
        run = paragraph.add_run(line)
        run.font.size = Pt(9)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_synthetic_schedule(output_path):
    output = Path(output_path).expanduser().resolve()
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = section.right_margin = Cm(1.5)
    section.top_margin = section.bottom_margin = Cm(1.5)

    title = document.add_paragraph("Synthetic College 2026 Schedule")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=5, cols=12)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "班级"
    table.rows[1].cells[0].text = "辅导员/电话"
    table.rows[0].cells[1].text = "人数"
    table.rows[1].cells[1].text = "人数"

    for index, day in enumerate(["星期一", "星期二", "星期三", "星期四", "星期五"]):
        start = 2 + index * 2
        cell = table.rows[0].cells[start]
        cell.text = day
        cell.merge(table.rows[0].cells[start + 1])
        set_cell_shading(cell, "D9E1F2")

    for index, period in enumerate(["1,2", "3,4"] * 5):
        table.rows[1].cells[index + 2].text = f"第{period}节"

    rows = [
        {
            "class": "测试专业2026级甲班",
            "advisor": "辅导员：测试辅导员甲",
            "phone": "SYNTHETIC-NO-PHONE",
            "student_count": "32",
            "slots": [
                ["测试课程A", "测试教师甲", "99号楼测试甲室", "全周"],
                ["测试课程B", "测试教师乙", "99号楼测试乙室", "单周上"],
                ["测试课程C", "测试教师丙", "99号楼测试丙室", "前八周"],
                [""],
                ["测试课程D", "测试教师甲", "99号楼测试甲室 或 99号楼测试乙室", "全周"],
                [""],
                ["测试课程E", "测试教师乙", "99号楼测试丁室", "后八周"],
                [""],
                [""],
                [""],
            ],
        },
        {
            "class": "测试专业2026级乙班",
            "advisor": "辅导员：测试辅导员乙",
            "phone": "SYNTHETIC-NO-PHONE",
            "student_count": "28",
            "slots": [
                ["测试课程B", "测试教师乙", "99号楼测试乙室", "双周上"],
                ["测试课程A", "测试教师甲", "99号楼测试甲室", "全周"],
                [""],
                ["测试课程C", "测试教师丙", "99号楼测试丙室", "后八周"],
                [""],
                ["测试课程D", "测试教师甲", "99号楼测试丁室", "前八周"],
                [""],
                [""],
                [""],
                [""],
            ],
        },
        {
            "class": "测试专业2026级丙班",
            "advisor": "辅导员：测试辅导员丙",
            "phone": "SYNTHETIC-NO-PHONE",
            "student_count": "30",
            "slots": [
                ["测试课程C", "测试教师丙", "99号楼测试丙室", "全周"],
                [""],
                ["测试课程E", "测试教师乙", "99号楼测试丁室", "单周上"],
                [""],
                ["测试课程A", "测试教师甲", "99号楼测试甲室", "双周上"],
                [""],
                [""],
                [""],
                [""],
                [""],
            ],
        },
    ]

    for row_index, data in enumerate(rows, start=2):
        row = table.rows[row_index]
        add_lines(row.cells[0], [data["class"], data["advisor"], data["phone"]])
        set_cell_shading(row.cells[0], "F2F2F2")
        row.cells[1].text = data["student_count"]
        for column_index, slot in enumerate(data["slots"], start=2):
            if slot and slot[0]:
                add_lines(row.cells[column_index], slot)

    document.save(output)
    print(f"Synthetic schedule generated: {output}")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output", help="Explicit output DOCX path")
    args = parser.parse_args()
    create_synthetic_schedule(args.output)


if __name__ == "__main__":
    main()
