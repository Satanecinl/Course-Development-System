#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
诊断脚本：只采样前 3 个班级的单元格原始文本
不做任何正则提取，仅打印包含换行符的脏数据样本
"""

from docx import Document
from pathlib import Path

# 文件路径
DOCX_PATH = r"C:\Users\Satanecinl\Desktop\Course Development System\2026年春季学期课程表(0420).docx"

def main():
    doc = Document(DOCX_PATH)
    print(f"文档共包含 {len(doc.tables)} 个表格\n")

    sample_count = 0
    max_samples = 10

    # 遍历所有表格
    for table_idx, table in enumerate(doc.tables):
        # 每个表格通常对应一个班级，取前 3 个有数据的表格
        if sample_count >= max_samples:
            break

        print(f"===== 表格 {table_idx + 1} =====")

        # 遍历所有行
        for row_idx, row in enumerate(table.rows):
            if sample_count >= max_samples:
                break

            # 遍历所有单元格
            for col_idx, cell in enumerate(row.cells):
                if sample_count >= max_samples:
                    break

                # 提取原始文本，保留换行符
                paragraphs = [p.text for p in cell.paragraphs]
                raw_text = '\n'.join(paragraphs)

                # 只打印包含换行符（多段落）的单元格
                if len(paragraphs) > 1 or '\x0b' in raw_text or '\n' in raw_text.strip():
                    # 进一步检查：排除纯空白和极短内容
                    stripped = raw_text.strip()
                    if len(stripped) > 2:
                        print(f"\n--- 样本 #{sample_count + 1} | 表格{table_idx+1} 行{row_idx} 列{col_idx} ---")
                        print(f"段落数: {len(paragraphs)}")
                        print(f"原始字符串（repr）:")
                        print(repr(raw_text))
                        print(f"--- 可读形式 ---")
                        print(raw_text)
                        print("-" * 40)
                        sample_count += 1

    print(f"\n共采集 {sample_count} 个脏数据样本")

if __name__ == '__main__':
    main()
