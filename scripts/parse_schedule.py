#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工程应用技术学院 Word 课程表解析脚本
使用 python-docx 和 pandas 解析复杂表格结构，输出结构化 JSON

表头结构特征：
- 第一行为星期（一、二、三、四、五），跨多列
- 第二行为节次（1,2 / 3,4 / 5,6 / 7,8 / 9,10）
- 第一列为班级信息（班级名、辅导员、电话）
- 合并单元格的空列继承左侧邻居的节次映射

单元格内容特征：
- 多行文本，每行包含：课程名、教师、教室、周次限制
- 教室格式：'11-322'（数字-数字）或 '林校305'（汉字+数字）
- 周次限制：'前八周'、'后八周'、'单周上'、'双周上'、'全周' 等
"""

import re
import sys
import json
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, asdict

# 添加 scripts 目录到 path，以便导入 parse_cell
sys.path.insert(0, str(Path(__file__).parent))

from parse_cell import parse_cell_content

try:
    from docx import Document
    from docx.table import Table
    from docx.oxml.ns import qn
except ImportError:
    raise ImportError("请安装 python-docx: pip install python-docx")

try:
    import pandas as pd
except ImportError:
    raise ImportError("请安装 pandas: pip install pandas")


# ===== 节次到 slot 的映射 =====
# slot_index: 1=1-2节, 2=3-4节, 3=5-6节, 4=7-8节, 5=9-10节, 6=11-12节(应排除), 7=中午
PERIOD_SLOT_MAP = {
    '1,2': 1, '1、2': 1, '1-2': 1, '1—2': 1,
    '3,4': 2, '3、4': 2, '3-4': 2, '3—4': 2,
    '5,6': 3, '5、6': 3, '5-6': 3, '5—6': 3,
    '7,8': 4, '7、8': 4, '7-8': 4, '7—8': 4,
    '9,10': 5, '9、10': 5, '9-10': 5, '9—10': 5,
    '11,12': 6, '11、12': 6, '11-12': 6, '11—12': 6,
}

# 伪记录过滤：表头行 / 节次文本 / 星期文本
HEADER_CLASS_NAMES = {'专业年级班', '人数', '教室'}
HEADER_STUDENT_RAW = {'人数'}
HEADER_COURSE_TEXTS = {
    '周一', '周二', '周三', '周四', '周五', '周六', '周日',
    '一', '二', '三', '四', '五', '六', '日',
    '1、2', '3、4', '5、6', '7、8', '9、10', '9.10',
    '1-2节', '3-4节', '5-6节', '7-8节', '9-10节',
    '1', '2', '3', '4', '5', '6', '7', '8', '9', '10',
    '中午 12点~13点30', '中午',
    '专业年级班', '人数', '教室',
}

VALID_SLOTS = {1, 2, 3, 4, 5, 7}  # 允许的有效 slot（排除 6=11-12节）


def is_valid_schedule_record(record: Dict[str, Any]) -> bool:
    """判断一条记录是否为有效课程记录（过滤表头行伪记录）。"""
    class_name = (record.get('class_info', {}).get('class_name') or '').strip()
    student_count_raw = (record.get('class_info', {}).get('student_count_raw') or '').strip()
    course = (record.get('course') or '').strip()

    if class_name in HEADER_CLASS_NAMES:
        return False
    if student_count_raw in HEADER_STUDENT_RAW:
        return False
    if course in HEADER_COURSE_TEXTS:
        return False
    return True


def deduplicate_records(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """对 parser 输出做保守去重：仅去除完全相同的 records，不合并真实不同时间课程。"""
    seen = set()
    result = []
    for r in records:
        # 构建去重 key
        class_name = r.get('class_info', {}).get('class_name', '') or ''
        key = '|'.join([
            class_name,
            r.get('course', '') or '',
            r.get('teacher', '') or '',
            r.get('room', '') or '',
            str(r.get('day_of_week', '')),
            r.get('time_slot', '') or '',
            r.get('week_type', '') or '',
            str(r.get('week_start', '')),
            str(r.get('week_end', '')),
            r.get('remark', '') or '',
        ])
        if key not in seen:
            seen.add(key)
            result.append(r)
    return result


def split_course_teacher_by_known_teacher(text: str, teacher_names: List[str]) -> Tuple[Optional[str], Optional[str]]:
    """
    用已知教师名单反向切分粘连的 课程名+教师名。
    例如 "无机化学张测试" → ("无机化学", "张测试")
    返回 (course, teacher) 或 (None, None) 表示无法切分。
    """
    if not text or not teacher_names:
        return None, None
    # 清理末尾括号残留
    text = re.sub(r'[\s（）()]+$', '', text.strip())
    if not text:
        return None, None
    # 按长度降序匹配（教师名单已排序，但再保险一次）
    for name in teacher_names:
        if text.endswith(name) and len(text) > len(name):
            course = text[:-len(name)].strip()
            if len(course) >= 2:
                return course, name
    return None, None


@dataclass
class ClassInfo:
    """班级信息"""
    class_name: str
    advisor_name: Optional[str] = None
    advisor_phone: Optional[str] = None
    student_count_raw: Optional[str] = None
    student_count: Optional[int] = None


@dataclass
class ParsedCell:
    """解析后的单元格数据"""
    course: Optional[str] = None
    teacher: Optional[str] = None
    room: Optional[str] = None
    week_constraints: Optional[str] = None
    week_start: int = 1
    week_end: int = 16
    week_type: str = "ALL"
    remark: Optional[str] = None


@dataclass
class ScheduleRecord:
    """最终输出的排课记录"""
    class_info: Dict[str, Any]
    teacher: Optional[str]
    course: Optional[str]
    room: Optional[str]
    day_of_week: int
    time_slot: str
    period_start: int
    period_end: int
    week_constraints: Optional[str]
    week_start: int
    week_end: int
    week_type: str
    remark: Optional[str] = None
    student_count_raw: Optional[str] = None
    student_count: Optional[int] = None


# ===== 正则表达式模式 =====
ROOM_PATTERNS = [
    re.compile(r'\d{1,2}-\d{3}[A-Z]?'),
    re.compile(r'[一-龥]+\d{3,4}[A-Z]?'),
    re.compile(r'[A-Z]-\d{3,4}'),
]

WEEK_CONSTRAINT_PATTERNS = {
    'FIRST_HALF': re.compile(r'前\s*八\s*周|前\s*8\s*周|1-8\s*周'),
    'SECOND_HALF': re.compile(r'后\s*八\s*周|后\s*8\s*周|9-16\s*周'),
    'ODD': re.compile(r'单\s*周\s*上|单\s*周'),
    'EVEN': re.compile(r'双\s*周\s*上|双\s*周'),
}

ALL_WEEKS_PATTERN = re.compile(r'全\s*周|全\s*学\s*期')
CUSTOM_WEEK_PATTERN = re.compile(r'(?:第)?(\d{1,2})\s*[-~]\s*(\d{1,2})\s*周')
TEACHER_PATTERN = re.compile(r'[一-龥]{2,4}')
PHONE_PATTERN = re.compile(r'1[3-9]\d{9}|\d{3,4}-\d{7,8}')
PHONE_MOBILE_RE = re.compile(r'1[3-9]\d{9}')
ADVISOR_PATTERN = re.compile(r'辅导[员師]\s*[:：]?\s*([一-龥]{2,4})')


def mask_pii(text: str) -> str:
    """对文本中的中国大陆手机号进行脱敏替换。"""
    if not text:
        return text
    return PHONE_MOBILE_RE.sub('[手机号已脱敏]', text)


def parse_student_count(raw: str) -> Optional[int]:
    """
    解析复合人数字符串为总和。
    支持：'25'、'18+14+2'、'21+15' 等。
    禁止使用 eval。无法解析时返回 None。
    """
    if not raw:
        return None
    text = raw.strip()
    if not text:
        return None
    # 正则校验：纯数字或数字+数字+...格式
    if not re.match(r'^\d+(?:\+\d+)*$', text):
        return None
    try:
        return sum(int(x) for x in text.split('+'))
    except (ValueError, TypeError):
        return None


# ========== 可选教师白名单 ==========

def load_teacher_names(txt_path: str) -> List[str]:
    """读取清洗后的教师白名单（已按长度降序排列）。"""
    path = Path(txt_path)
    if not path.exists():
        print("警告：已配置的教师白名单不可用，将使用无白名单模式")
        return []
    with open(path, "r", encoding="utf-8") as f:
        names = [line.strip() for line in f if line.strip()]
    return names


def build_teacher_regex(teacher_names: List[str]) -> Optional[re.Pattern]:
    """将教师姓名列表编译为正则（最长优先，避免部分匹配）。"""
    if not teacher_names:
        return None
    # 白名单已按长度降序排列
    escaped = [re.escape(name) for name in teacher_names]
    return re.compile('|'.join(escaped))


# ========== 文本提取 ==========

def extract_text_from_cell(cell) -> str:
    """提取单元格中的所有文本，保留换行"""
    paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return '\n'.join(paragraphs)


def parse_cell_text(
    text: str,
    teacher_regex: Optional[re.Pattern] = None,
) -> List[ParsedCell]:
    """
    解析单元格内的文本（使用 parse_cell_content 引擎）。
    若提供 teacher_regex，则在 parse_cell_content 内部进行白名单权威匹配。
    """
    if not text or not text.strip():
        return []

    records = parse_cell_content(text, teacher_regex)

    results = []
    for r in records:
        pc = ParsedCell()
        pc.course = r.get("course_name")
        pc.teacher = r.get("teacher")
        pc.remark = r.get("remark")

        classroom = r.get("classroom")
        if isinstance(classroom, list):
            pc.room = " 或 ".join(classroom)
        elif isinstance(classroom, str):
            pc.room = classroom

        weeks = r.get("weeks", "")
        if weeks:
            pc.week_constraints = weeks
            if re.search(r'前\s*八\s*周|前\s*8\s*周', weeks):
                pc.week_type = "FIRST_HALF"
                pc.week_start, pc.week_end = 1, 8
            elif re.search(r'后\s*八\s*周|后\s*8\s*周', weeks):
                pc.week_type = "SECOND_HALF"
                pc.week_start, pc.week_end = 9, 16
            elif re.search(r'单\s*周', weeks):
                pc.week_type = "ODD"
            elif re.search(r'双\s*周', weeks):
                pc.week_type = "EVEN"
            else:
                m = re.match(r'(\d+)-(\d+)周', weeks)
                if m:
                    pc.week_type = "CUSTOM"
                    pc.week_start = int(m.group(1))
                    pc.week_end = int(m.group(2))

        results.append(pc)

    return results


def parse_class_info(text: str) -> ClassInfo:
    """解析第一列的班级信息，并对手机号进行 PII 脱敏。"""
    lines = [line.strip() for line in text.split('\n') if line.strip()]

    class_name = None
    advisor_name = None
    advisor_phone = None

    for line in lines:
        advisor_match = ADVISOR_PATTERN.search(line)
        if advisor_match:
            advisor_name = advisor_match.group(1)
            continue

        # 先在原始文本上检测是否存在手机号
        phone_match = PHONE_PATTERN.search(line)
        if phone_match:
            advisor_phone = '[手机号已脱敏]'
            continue

        if '班' in line and not class_name:
            class_name = line
            continue

        if not class_name and not advisor_name and not advisor_phone:
            class_name = line

    # 对 class_name 做脱敏（防止号码出现在班级名中）
    if class_name:
        class_name = mask_pii(class_name)

    return ClassInfo(
        class_name=class_name or mask_pii(text.strip()),
        advisor_name=advisor_name,
        advisor_phone=advisor_phone
    )


# ========== 表头解析（核心修复 1） ==========

def _normalize_period(text: str) -> Optional[int]:
    """
    将节次文本归一化为 slot 编号。
    返回 None 表示无法识别；返回 6 表示 11-12 节（需被过滤）。
    """
    if not text:
        return None

    # 去掉空格和"节"字
    normalized = text.replace(' ', '').replace('节', '')

    # 直接查表
    slot = PERIOD_SLOT_MAP.get(normalized)
    if slot is not None:
        return slot

    # 正则匹配 "数字[,、/-]数字"
    m = re.search(r'(\d{1,2})\s*[,，、/-—]\s*(\d{1,2})', text)
    if m:
        key = f"{m.group(1)},{m.group(2)}"
        return PERIOD_SLOT_MAP.get(key)

    # 单数字：映射到对应区间
    single_match = re.search(r'(\d{1,2})', text)
    if single_match:
        p = int(single_match.group(1))
        if p <= 2:
            return 1
        elif p <= 4:
            return 2
        elif p <= 6:
            return 3
        elif p <= 8:
            return 4
        elif p <= 10:
            return 5
        elif p <= 12:
            return 6

    # 中午
    if '中午' in text or '午休' in text:
        return 7

    return None


def _get_cell_xml_info(cell) -> Dict[str, Any]:
    """提取单元格的 XML 属性：gridSpan（水平合并列数）和 vMerge（垂直合并）"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    info = {'gridSpan': 1, 'vMerge': None}
    if tcPr is not None:
        gridSpan_elem = tcPr.find(qn('w:gridSpan'))
        if gridSpan_elem is not None:
            info['gridSpan'] = int(gridSpan_elem.get(qn('w:val'), '1'))
        vMerge_elem = tcPr.find(qn('w:vMerge'))
        if vMerge_elem is not None:
            info['vMerge'] = vMerge_elem.get(qn('w:val'), 'continue')
    return info


def _is_day_header_row(texts: List[str]) -> bool:
    """判断一行是否为星期行：至少 5 个非空 cell 包含星期字符。"""
    day_count = sum(
        1 for t in texts
        if t and any(d in t for d in '一二三四五六日')
    )
    return day_count >= 5


def _is_period_header_row(texts: List[str]) -> bool:
    """判断一行是否为节次行：至少 5 个非空 cell 包含节次数字模式。"""
    period_count = sum(
        1 for t in texts
        if t and re.search(r'\d{1,2}\s*[,，、./-—]\s*\d{1,2}', t)
    )
    return period_count >= 5


def parse_header_rows(table: Table, table_idx: int = 0, verbose: bool = False) -> Tuple[int, List[Dict[str, Any]]]:
    """
    解析表头，确定数据起始行和列映射。
    关键修复：空列继承左侧邻居的 slot；仅保留 VALID_SLOTS（过滤 11-12 节）。
    返回：(数据起始行索引, 列映射列表)
    """
    header_rows = []
    for i, row in enumerate(table.rows):
        cells_text = [extract_text_from_cell(cell).strip() for cell in row.cells]
        header_rows.append((i, cells_text))

    day_row_idx = None
    period_row_idx = None

    # ── 策略 1：优先信任前两行的标准表头结构 ──
    if len(header_rows) >= 2:
        if _is_day_header_row(header_rows[0][1]) and _is_period_header_row(header_rows[1][1]):
            day_row_idx = 0
            period_row_idx = 1

    # ── 策略 2：回退到模糊检测（仅当策略 1 失败时）──
    if day_row_idx is None:
        for idx, texts in header_rows[:4]:
            has_days = any(
                ('周' in t and any(d in t for d in '一二三四五六日')) or
                any(d in t for d in '一二三四五六日')
                for t in texts if t
            )
            has_periods = any('节' in t and re.search(r'\d{1,2}', t) for t in texts if t)

            if has_days and not has_periods:
                day_row_idx = idx
            elif has_periods and day_row_idx is not None:
                period_row_idx = idx
                break

        if day_row_idx is None:
            for idx, texts in header_rows[:3]:
                if any(any(d in t for d in '一二三四五六日') for t in texts if t):
                    day_row_idx = idx
                    break

        if period_row_idx is None and day_row_idx is not None:
            next_idx = day_row_idx + 1
            if next_idx < len(header_rows):
                period_row_idx = next_idx

    data_start_row = (period_row_idx or day_row_idx or 0) + 1

    column_map = []
    day_names = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '日': 7}

    if day_row_idx is not None and period_row_idx is not None:
        day_texts = header_rows[day_row_idx][1]
        period_texts = header_rows[period_row_idx][1]

        current_day = None
        last_slot = None

        for col_idx, (day_text, period_text) in enumerate(zip(day_texts, period_texts)):
            if col_idx == 0:
                continue  # 跳过班级列

            # 更新当前星期
            if day_text:
                for day_char, day_num in day_names.items():
                    if day_char in day_text:
                        current_day = day_num
                        break

            # 解析节次
            slot = _normalize_period(period_text)

            # 空列继承左侧邻居
            if slot is None and last_slot is not None:
                slot = last_slot

            # 仅保留有效 slot（过滤 6=11-12 节）
            if slot is not None and slot in VALID_SLOTS and current_day is not None:
                period_map = {
                    1: (1, 2), 2: (3, 4), 3: (5, 6),
                    4: (7, 8), 5: (9, 10), 7: (0, 0),
                }
                period_start, period_end = period_map.get(slot, (0, 0))
                time_slot = f"{period_start},{period_end}" if slot != 7 else "中午"

                column_map.append({
                    'col_idx': col_idx,
                    'day_of_week': current_day,
                    'time_slot': time_slot,
                    'period_start': period_start,
                    'period_end': period_end,
                    'slot_index': slot,
                })
                last_slot = slot
            elif slot is not None:
                #  slot 存在但不在 VALID_SLOTS（如 6=11-12），更新 last_slot 但不加入映射
                last_slot = slot

    # ========== 诊断输出 ==========
    if verbose:
        print(f"\n{'='*60}")
        print(f"【诊断】表格 {table_idx + 1} 表头解析")
        print(f"{'='*60}")
        print(f"  星期行索引: {day_row_idx}")
        print(f"  节次行索引: {period_row_idx}")
        print(f"  数据起始行: {data_start_row}")
        print(f"\n  星期行原始文本 (row.cells):")
        if day_row_idx is not None:
            for ci, t in enumerate(header_rows[day_row_idx][1]):
                print(f"    col[{ci:2d}]: '{t}'")
        print(f"\n  节次行原始文本 (row.cells):")
        if period_row_idx is not None:
            for ci, t in enumerate(header_rows[period_row_idx][1]):
                print(f"    col[{ci:2d}]: '{t}'")
        print(f"\n  列映射字典 ({len(column_map)} 列):")
        for cm in column_map:
            print(f"    col_idx={cm['col_idx']:2d} -> 周{cm['day_of_week']} {cm['time_slot']} (slot={cm['slot_index']})")

    return data_start_row, column_map


# ========== 表格解析 ==========

def parse_schedule_table(
    table: Table,
    teacher_regex: Optional[re.Pattern] = None,
    table_idx: int = 0,
    verbose: bool = False,
) -> List[ScheduleRecord]:
    """解析单个课程表"""
    records = []
    data_start_row, column_map = parse_header_rows(table, table_idx, verbose)

    if not column_map:
        print("警告：未能解析表头结构")
        return records

    # 检测"人数"列：扫描第一行（星期行）的 header 文本
    student_count_col: Optional[int] = None
    if len(table.rows) > 0:
        header_cells = table.rows[0].cells
        for ci, cell in enumerate(header_cells):
            text = extract_text_from_cell(cell).strip()
            if '人数' in text:
                student_count_col = ci
                break

    if verbose and student_count_col is not None:
        print(f"\n  【人数列检测】col_idx={student_count_col}")

    # 诊断：抽取一个样本班级，打印其原始数据行
    sample_class_rows = []

    for row_idx in range(data_start_row, len(table.rows)):
        row = table.rows[row_idx]
        cells = row.cells

        if len(cells) == 0:
            continue

        # 第一列：班级信息
        class_text = extract_text_from_cell(cells[0])
        if not class_text.strip():
            continue

        class_info = parse_class_info(class_text)

        # 提取人数列（如果存在）
        if student_count_col is not None and student_count_col < len(cells):
            sc_raw = extract_text_from_cell(cells[student_count_col]).strip()
            if sc_raw:
                class_info.student_count_raw = sc_raw
                class_info.student_count = parse_student_count(sc_raw)

        # 收集样本数据（前 3 个非空班级行）
        if verbose and len(sample_class_rows) < 3:
            raw_cells = []
            for col_info in column_map:
                col_idx = col_info['col_idx']
                if col_idx >= len(cells):
                    raw_cells.append("")
                else:
                    raw_cells.append(extract_text_from_cell(cells[col_idx]).replace('\n', '\\n'))
            sample_class_rows.append({
                'class_name': class_info.class_name,
                'row_idx': row_idx,
                'raw_cells': raw_cells,
                'column_map': column_map,
            })

        # 遍历各时间段列
        for col_info in column_map:
            col_idx = col_info['col_idx']
            if col_idx >= len(cells):
                continue

            cell_text = extract_text_from_cell(cells[col_idx])
            if not cell_text.strip():
                continue

            parsed_cells = parse_cell_text(cell_text, teacher_regex)

            for parsed in parsed_cells:
                record = ScheduleRecord(
                    class_info=asdict(class_info),
                    teacher=parsed.teacher,
                    course=parsed.course,
                    room=parsed.room,
                    day_of_week=col_info['day_of_week'],
                    time_slot=col_info['time_slot'],
                    period_start=col_info['period_start'],
                    period_end=col_info['period_end'],
                    week_constraints=parsed.week_constraints,
                    week_start=parsed.week_start,
                    week_end=parsed.week_end,
                    week_type=parsed.week_type,
                    remark=parsed.remark,
                    student_count_raw=class_info.student_count_raw,
                    student_count=class_info.student_count,
                )
                records.append(record)

    # 打印样本班级原始数据
    if verbose and sample_class_rows:
        print(f"\n  【诊断】样本班级原始数据行（前 {len(sample_class_rows)} 个）:")
        for sr in sample_class_rows:
            print(f"\n  班级: {sr['class_name']} (row {sr['row_idx']})")
            for ci, cell_text in enumerate(sr['raw_cells']):
                cm = sr['column_map'][ci]
                print(f"    col[{cm['col_idx']:2d}] (周{cm['day_of_week']} {cm['time_slot']:>5s}): '{cell_text[:60]}{'...' if len(cell_text) > 60 else ''}'")

    return records


def parse_word_schedule(
    docx_path: str,
    output_path: Optional[str] = None,
    teachers_path: Optional[str] = None,
    verbose: bool = False,
) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    """
    解析 Word 课程表文件。
    若提供 teachers_path，则加载教师白名单进行权威匹配。
    """
    teacher_regex = None
    teacher_names: List[str] = []
    if teachers_path:
        teacher_names = load_teacher_names(teachers_path)
        teacher_regex = build_teacher_regex(teacher_names)
        if teacher_regex:
            print(f"教师白名单: 已启用 ({len(teacher_names)} 条)")
    else:
        print("教师白名单: 未配置，使用无白名单模式")

    doc = Document(docx_path)

    all_records = []
    total_cells = 0
    multi_split_cells = 0
    class_names = set()

    for table_idx, table in enumerate(doc.tables):
        print(f"正在解析第 {table_idx + 1} 个表格...")
        records = parse_schedule_table(table, teacher_regex, table_idx, verbose=verbose)
        all_records.extend(records)
        print(f"  提取到 {len(records)} 条记录")

        # 统计
        data_start_row, column_map = parse_header_rows(table, table_idx, verbose=verbose)
        for row_idx in range(data_start_row, len(table.rows)):
            row = table.rows[row_idx]
            cells = row.cells
            if len(cells) == 0:
                continue
            class_text = extract_text_from_cell(cells[0])
            if class_text.strip():
                class_info = parse_class_info(class_text)
                if class_info.class_name:
                    class_names.add(class_info.class_name)

            for col_info in column_map:
                col_idx = col_info['col_idx']
                if col_idx >= len(cells):
                    continue
                cell_text = extract_text_from_cell(cells[col_idx])
                if not cell_text.strip():
                    continue
                total_cells += 1
                parsed = parse_cell_text(cell_text, teacher_regex)
                if len(parsed) >= 2:
                    multi_split_cells += 1

    # 过滤伪记录（表头行、节次文本等）
    valid_records = [r for r in all_records if is_valid_schedule_record(asdict(r))]

    # 保守去重：仅去除完全相同的 records
    valid_records_before_dedup = len(valid_records)
    valid_records = deduplicate_records([asdict(r) for r in valid_records])
    # 将 dict 转回 ScheduleRecord
    valid_records = [ScheduleRecord(**r) for r in valid_records]
    dedup_count = valid_records_before_dedup - len(valid_records)
    if dedup_count > 0:
        print(f"去重 records: {dedup_count} 条")

    # 修复教师名粘连：用已知教师名单反向切分
    if teacher_names:
        for r in valid_records:
            if not r.teacher and r.course:
                course, teacher = split_course_teacher_by_known_teacher(r.course, teacher_names)
                if teacher:
                    r.course = course
                    r.teacher = teacher

    result = [asdict(r) for r in valid_records]

    filtered_count = len(all_records) - len(valid_records)
    if filtered_count > 0:
        print(f"过滤伪记录: {filtered_count} 条")

    stats = {
        "class_count": len(class_names),
        "total_records": len(valid_records),
        "total_nonempty_cells": total_cells,
        "multi_split_cells": multi_split_cells,
    }

    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"\n结果已保存到: {output_path}")

    return result, stats


# ========== DataFrame / 主入口 ==========

def to_dataframe(records: List[Dict[str, Any]]) -> pd.DataFrame:
    """将解析结果转换为 pandas DataFrame"""
    if not records:
        return pd.DataFrame()

    df = pd.DataFrame(records)

    if 'class_info' in df.columns:
        class_df = pd.json_normalize(df['class_info'])
        df = df.drop('class_info', axis=1)
        df = pd.concat([class_df, df], axis=1)

    preferred_order = [
        'class_name', 'advisor_name', 'advisor_phone',
        'course', 'teacher', 'room',
        'day_of_week', 'time_slot',
        'week_start', 'week_end', 'week_type', 'week_constraints', 'remark'
    ]

    existing_cols = [c for c in preferred_order if c in df.columns]
    other_cols = [c for c in df.columns if c not in preferred_order]
    df = df[existing_cols + other_cols]

    return df


def main():
    parser = argparse.ArgumentParser(description='解析工程应用技术学院 Word 课程表')
    parser.add_argument('input', help='输入 Word 文件路径')
    parser.add_argument('-o', '--output', help='输出 JSON 文件路径')
    parser.add_argument('-c', '--csv', help='输出 CSV 文件路径')
    parser.add_argument('-v', '--verbose', action='store_true', help='显示详细解析信息')
    parser.add_argument(
        '--teachers',
        default=None,
        help='可选教师白名单文件路径；未提供时使用无白名单模式',
    )

    args = parser.parse_args()

    if not Path(args.input).exists():
        print(f"错误：文件不存在: {args.input}")
        return 1

    records, stats = parse_word_schedule(args.input, args.output, args.teachers, verbose=args.verbose)

    print(f"\n========== 统计信息 ==========")
    print(f"成功解析班级数:     {stats['class_count']}")
    print(f"非空单元格总数:     {stats['total_nonempty_cells']}")
    print(f"独立课程记录总数:   {stats['total_records']}")
    print(f"触发多课拆分单元格: {stats['multi_split_cells']}")
    print(f"===============================")

    if args.verbose and records:
        print("\n前 3 条记录示例:")
        for r in records[:3]:
            print(json.dumps(r, ensure_ascii=False, indent=2))

    if args.csv:
        df = to_dataframe(records)
        df.to_csv(args.csv, index=False, encoding='utf-8-sig')
        print(f"CSV 已保存到: {args.csv}")

    return 0


if __name__ == '__main__':
    exit(main())
